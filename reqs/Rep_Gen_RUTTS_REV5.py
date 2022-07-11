#!/usr/bin/env python
# coding: utf-8

# In[ ]:


get_ipython().run_line_magic('matplotlib', 'inline')
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import tkinter as tk
from tkinter import *
from tkinter import simpledialog
from tkinter import messagebox
from tkinter import filedialog
from tkinter import ttk
#import ntpath
import datetime
import docx
import pandas as pd
from openpyxl import load_workbook
import matplotlib.pyplot as plt
import matplotlib.ticker as ticker
from ipywidgets import interactive
import numpy as np
import ipywidgets as widgets
from IPython.display import display, clear_output
import voila


# In[ ]:


class Report_template:
    device_des = ''
    freq = ''
    sup_vol = ''
    output_type = ''
    dev_code = ''
    rev_day = ''
    rev_month = ''
    rev_year = ''
    n = ''
    so_num = ''
    name = ''
    test_cond = ''
    test1_cond = ''
    test2_cond = ''
    test_set_1 = ''
    test1_set_1 = ''
    test2_set_1 = ''
    fvt_spec = 0
    temp_range = [0]
    plot_num = 0
    pn_spec = [0]
    save_path = ''
    ebt_num = 0
    pn_fl_type = 'csv'
    
def Merge_context(contxt1, contxt2):
    contxt_full = {**contxt1, **contxt2}
    return contxt_full

temp = Report_template()
cols_2 = []

def get_Table_1(template_path):
    messagebox.showinfo('Info','Please select the Electical Parameters file(.txt).')
    table_path = filedialog.askopenfilename()
    # ebt_num 
    #rows_2_read =[*range(6,10+int(temp.ebt_num),1)]
    rows_2_read =[*range(10,10+int(temp.ebt_num),1)]
    rows_2_read_temp =[*range(6,9,1)]
    read_file = pd.read_csv (table_path, skiprows= lambda x:x not in rows_2_read, delim_whitespace=False)
    read_file_temp = pd.read_csv (table_path, skiprows= lambda x:x not in rows_2_read_temp, delim_whitespace = False)
    
    read_file.to_csv ('C:/Xice/python_excel/BenchMeasurements.csv', index=False)
    read_file_temp.to_csv ('C:/Xice/python_excel/BenchMeasurements_temp.csv', index=False)

    fl = pd.read_csv ('C:/Xice/python_excel/BenchMeasurements.csv', index_col = False)
    fl1 = pd.read_csv ('C:/Xice/python_excel/BenchMeasurements_temp.csv', index_col = False)
    #print(fl)
    fl = fl.to_string(index=False)
    fl1 = fl1.to_string(index=False)
    #fl=fl.replace("  ", "      ")
    #fl=fl.replace(" ", "      ")
    #fl=fl.replace("      "*6, "      ")
    #fl=fl.replace("            ", "      ")
    #fl=fl.replace("NaN", "")
    #print(fl)
    fl1=fl1.replace("                  ", "                            ")
    fl1=fl1.replace("  ", "    ")
    fl1=fl1.replace("             ", "       ")
    fl1=fl1.replace("                                ", "                                     ")
    fl1=fl1.replace("          ", "                ")
    fl1=fl1.replace("         ", "        ")
    fl1=fl1.replace("                                                 ", "                                      ")
    fl=fl.replace(" ", "  ")

    mydoc = docx.Document(template_path)

    # t=mydoc.add_table(fl.shape[0]+1,fl.shape[1])
    # for j in range(fl.shape[-1]):
        #     t.cell(0,j).text=fl.columns[j]
    
    # for i in range(fl.shape[0]):
        #     for j in range(fl.shape[-1]):
            #         t.cell(i+1,j).text = str(fl.values[i,j])

    # mydoc.save("C:/Xice/python_excel/template_test.docx")
    # doc = docx.Document() 

    #doc.save('C:/Xice/python_excel/table.docx')
    for para in mydoc.paragraphs:
        if 'TABLE_1' in para.text:
            para.text = fl1
            #print(para.text)
        if 'TABLE_2' in para.text:
            para.text = fl
      
    # mydoc.save("C:/Xice/python_excel/template_test.docx")
    mydoc.save(template_path)
    
def PN_gen_csv():
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo('Info','Please select your raw Phase Noise file. ')
    # plot_num
    d = {}
    for num in range(temp.plot_num):
        file_path = filedialog.askopenfilename()
        print(file_path)
        # unit_no = file_path[-12:-4:1] # To present Unit marked No#
        unit_no = "Unit" + str(num+1) # To present Unit No# in sequence
        print(unit_no)
        d[num]=pd.read_csv(file_path)
        d[num].rename(columns = {'Phase Noise (dBc/Hz)' : 'Phase Noise (dBc/Hz) - '+ unit_no}, inplace = True)
        print(d[num])
        #d[num].iloc[:,1] = d[num].iloc[:,1].rolling(15).mean() # Calculating Simple Moving Average
        d[num].iloc[:,1] = d[num].iloc[:,1].ewm(span = 30).mean() # Calculating Exponential Moving Average
        #x = list(d[num]['Offset Frequency (Hz)'])
        #y = list(d[num]['Phase Noise (dBc/Hz)'])
        lst = [1, 10, 100, 1000, 10000, 100000, 1000000] 
        #lst = [10, 100, 1000, 10000, 100000, 1000000] # as per SPEC
        df_spec = pd.DataFrame()
        df_spec['Offset'] = lst
        df_spec['PN typical'] = temp.pn_spec

# df = pd.concat([d[0], d[1]])
# print(df)
    
# # x = list(df.iloc[0:,0])
# # y = list(df.iloc[0:,1])
# print(x)
# print(y)
    for i in range(temp.plot_num):
        if i == 0:
            ax = df_spec.plot(x = 'Offset', y = 'PN typical', kind ='line', linestyle='dashed', logx = True, figsize = (16, 8), grid = True)
            d[i].plot(ax = ax, x = 'Offset Frequency (Hz)', y = d[i].columns[1], kind ='line', ylabel = 'Phase Noise(dBc/Hz)', title='Phase Noise', logx = True, figsize = (16, 8), grid = True)
        if(i > 0):
                d[i].plot(ax = ax, x = 'Offset Frequency (Hz)', y = d[i].columns[1], kind ='line', ylabel = 'Phase Noise(dBc/Hz)', title='Phase Noise', logx = True, figsize = (16, 8), grid = True)
    fig = ax.get_figure()
    fig.savefig("PN_output.png")
    
def PN_gen_txt():
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo('Info','Please select your raw Phase Noise file. ')
    # plot_num
    d = {}
    for num in range(temp.plot_num):
        file_path = filedialog.askopenfilename()
        print(file_path)
        #unit_no = file_path[-12:-4:1] # To present Unit marked No#
        unit_no = "Unit" + str(num+1) # To present Unit No# in sequence
        print(unit_no)
        read_file = pd.read_csv (file_path, delim_whitespace = True)
        read_file.to_csv(str(unit_no)+'.csv', index=None)
        file_path = str(unit_no)+'.csv'
        d[num]=pd.read_csv(file_path)
        d[num] = d[num][:-1]
        d[num].columns = ['Offset Frequency (Hz)', 'Phase Noise (dBc/Hz) - '+ unit_no]
        print(d[num])
        #d[num].iloc[:,1] = d[num].iloc[:,1].rolling(15).mean() # Calculating Simple Moving Average
        d[num].iloc[:,1] = d[num].iloc[:,1].ewm(span = 30).mean() # Calculating Exponential Moving Average
        #print(d[num])
        #x = list(d[num]['Offset Frequency (Hz)'])
        #y = list(d[num]['Phase Noise (dBc/Hz)'])
        lst = [1, 10, 100, 1000, 10000, 100000, 1000000]
        #lst = [10, 100, 1000, 10000, 100000, 1000000] # as per SPEC
        df_spec = pd.DataFrame()
        df_spec['Offset'] = lst
        df_spec['PN typical'] = temp.pn_spec

# df = pd.concat([d[0], d[1]])
# print(df)
    
# # x = list(df.iloc[0:,0])
# # y = list(df.iloc[0:,1])
# print(x)
# print(y)
    for i in range(temp.plot_num):
        if i == 0:
            ax = df_spec.plot(x = 'Offset', y = 'PN typical', kind ='line', linestyle='dashed', logx = True, figsize = (16, 8), grid = True)
            d[i].plot(ax = ax, x = 'Offset Frequency (Hz)', y = d[i].columns[1], kind ='line', ylabel = 'Phase Noise(dBc/Hz)', title='Phase Noise', logx = True, figsize = (16, 8), grid = True)
        if(i > 0):
                d[i].plot(ax = ax, x = 'Offset Frequency (Hz)', y = d[i].columns[1], kind ='line', ylabel = 'Phase Noise(dBc/Hz)', title='Phase Noise', logx = True, figsize = (16, 8), grid = True)
    fig = ax.get_figure()
    fig.savefig("PN_output.png")
    
    
def FvT_gen():
    root = tk.Tk()
    root.withdraw()
    messagebox.showinfo('Info','Please select your raw FvT data(.xlsx). ')
    file_path = filedialog.askopenfilename()
    print(file_path)
    
    if(file_path[-3:] == 'csv'):
        read_file = pd.read_csv(file_path)
        file_path = 'fvt_csv_cvted' + '.xlsx'
        print(file_path)
        read_file.to_excel(file_path,index = None)
        
    wb = load_workbook(filename= file_path, read_only=True)
    sheet = wb.sheetnames
    for i in sheet:
        print(i)
    ws = wb[sheet[0]]
    data_rows = []
    for row in ws.iter_rows(ws.min_row, ws.max_row):
    # for row in ws['A1':'E5727']:
        data_cols = []
        for cell in row:
            data_cols.append(cell.value)
        data_rows.append(data_cols)
    print(ws.max_row)
    df = pd.DataFrame(data_rows)
    #print(df)
    row_cnt = df.shape[0]
    
    #print(df.drop(df.columns[range(9)], axis = 1))
    df = df.drop(df.columns[range(9)], axis = 1)
    print(df.T)
    df = df.T # transpose dataframe from col to row
    cols = [1]
    FvT_gen.df2 = df[df.columns[cols]]
    #print(FvT_gen.df2)
    global cols_2
    for i in range(2,row_cnt,3): #get needed data rows/col
        cols_2.append(i)
    #print(cols_2)
    df = df[df.columns[cols_2]]
    #print(df)
    FvT_gen.result = pd.concat([FvT_gen.df2, df],axis=1)
    #print(FvT_gen.result)
    
    FvT_gen.result = FvT_gen.result.T.reset_index(drop=True).T ##reset columns index starts from 0
    FvT_gen.result = FvT_gen.result.astype('float64') ## rounding the dataframe decimal place
    FvT_gen.result.round(1)
    print(FvT_gen.result)
    
    ## res created for later manipulation
    FvT_gen.res = FvT_gen.df2.T.reset_index(drop=True).T
    FvT_gen.res = FvT_gen.res.astype('float64')
    FvT_gen.res.round(1)
    print(FvT_gen.res)
    
def clicked_start():

    
    # temp.device_des
    # temp.freq
    # temp.sup_vol
    # temp.output_type
    # temp.dev_code
    # temp.rev_day
    # temp.rev_month
    # temp.rev_year
    # temp.n
    # temp.so_num
    # temp.name
    # temp.test_cond
    # temp.test1_cond
    # temp.test2_cond
    # temp.test_set_1
    # temp.test1_set_1
    # temp.test2_set_1
    # temp.first_row
    # temp.rows_perUnit
    # temp.interval
    # temp.df_num
    # temp.plot_num
    # temp.save_path
    # temp.ebt_num

    
    temp.device_des = str(txt_1.get())
    temp.freq = str(txt_2.get())
    temp.sup_vol = str(txt_3.get())
    temp.output_type = str(txt_4.get())
    temp.dev_code = str(txt_5.get())
    temp.rev_day = str(txt_6.get())
    temp.rev_month = str(txt_7.get())
    temp.rev_year = str(txt_8.get())
    temp.n = str(txt_9.get())
    temp.so_num = str(txt_10.get())
    temp.name = str(txt_11.get())
    temp.ebt_num = int(txt_12.get())
    temp.test_cond = txt_13.get("1.0", "end - 1 chars")
    temp.test_set_1 = str(txt_14.get())
    temp.fvt_spec = float(txt_15.get())
    temp.temp_range = txt_16.get()
    temp.test1_cond = txt_17.get("1.0", "end - 1 chars")
    temp.test1_set_1 = str(txt_18.get())
    temp.plot_num = int(txt_19.get())
    temp.pn_spec = txt_21.get()
    temp.test2_cond = txt_22.get("1.0", "end - 1 chars")
    temp.test2_set_1 = str(txt_23.get())
    temp.save_path = str(txt_24.get())+'.docx'
    temp.pn_fl_type = var.get()
    
    try:
        temp.temp_range = temp.temp_range.split(' ')
        temp.temp_range = list(map(int, temp.temp_range))
        
    except:
        print("Wrong Operating temperature range!")
    
    
    try:
        temp.pn_spec = temp.pn_spec.split(' ')
        temp.pn_spec = list(map(int, temp.pn_spec))
        
    except:
        print("Wrong PN limit!")
        
    
    

    if(len(temp.save_path) <= 30):
        success = True
    else:
        msg_temp = tk.Tk()
        msg_temp.withdraw()
        messagebox.showerror("Error","File name too long")
        raise Exception("File name too long!")
        success = False


    if success:
        root.destroy()
        root.quit()
        #main()
        
        #sys.exit()


# GUI
# **************************************************************    

root = Tk()
root.title('Report Settings')
root.geometry('870x860')


space_0 = Label(root, text=" ", font='arial 12 bold')
space_0.grid(column=1, row=1)
lbl_0 = Label(root, text="Sample Order Info - ", font='arial 12 bold')
lbl_0.grid(column=1, row=2) 
##----------------------------------------------------------------------------

var = IntVar()
var.set(temp.device_des)



labl_1 = Label(root, text="Code/ FG Device Description")
labl_1.grid(column=1, row=4, sticky=E)
txt_1 = Entry(root, width=20, textvariable=var) 
txt_1.grid(column=2, row=4)
lbl_11 = Label(root, text=" ")
lbl_11.grid(column=3, row=4, sticky=W)

##----------------------------------------------------------------------------
var = IntVar()
var.set(temp.freq)

labl_2 = Label(root, text="Frequency")
labl_2.grid(column=4, row=4, sticky=E)
txt_2 = Entry(root, width=20, textvariable=var) 
txt_2.grid(column=5, row=4)
lbl_21 = Label(root, text="MHz")
lbl_21.grid(column=6, row=4, sticky=W)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.sup_vol)



labl_3 = Label(root, text="Supply Voltage")
labl_3.grid(column=1, row=5, sticky=E)
txt_3 = Entry(root, width=20, textvariable=var) 
txt_3.grid(column=2, row=5)
lbl_31 = Label(root, text="V")
lbl_31.grid(column=3, row=5, sticky=W)

##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.output_type)



labl_4 = Label(root, text="Output Type")
labl_4.grid(column=4, row=5, sticky=E)
txt_4 = Entry(root, width=20, textvariable=var) 
txt_4.grid(column=5, row=5)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.dev_code)



labl_5 = Label(root, text="Device Code")
labl_5.grid(column=1, row=6, sticky=E)
txt_5 = Entry(root, width=20, textvariable=var) 
txt_5.grid(column=2, row=6)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.rev_day)



labl_6 = Label(root, text="Spec Issue Rev date")
labl_6.grid(column=1, row=7, sticky=E)
txt_6 = Entry(root, width=20, textvariable=var) 
txt_6.grid(column=2, row=7)
lbl_61 = Label(root, text="Day")
lbl_61.grid(column=3, row=7, sticky=W)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.rev_month)



#labl_7 = Label(root, text="Spec Issue Rev date")
#labl_7.grid(column=1, row=10, sticky=E)
txt_7 = Entry(root, width=20, textvariable=var) 
txt_7.grid(column=2, row=8)
lbl_71 = Label(root, text="Month")
lbl_71.grid(column=3, row=8, sticky=W)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.rev_year)



#labl_8 = Label(root, text="Spec Issue Rev date")
#labl_8.grid(column=1, row=11, sticky=E)
txt_8 = Entry(root, width=20, textvariable=var) 
txt_8.grid(column=2, row=9)
lbl_81 = Label(root, text="Year")
lbl_81.grid(column=3, row=9, sticky=W)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.n)



labl_9 = Label(root, text="Sample Quantity")
labl_9.grid(column=4, row=7, sticky=E)
txt_9 = Entry(root, width=20, textvariable=var) 
txt_9.grid(column=5, row=7)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.so_num)



labl_10 = Label(root, text="SO#")
labl_10.grid(column=4, row=8, sticky=E)
txt_10 = Entry(root, width=20, textvariable=var) 
txt_10.grid(column=5, row=8)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.name)



labl_11 = Label(root, text="Author")
labl_11.grid(column=4, row=9, sticky=E)
txt_11 = Entry(root, width=20, textvariable=var) 
txt_11.grid(column=5, row=9)
space_1 = Label(root, text=" ", font='arial 12 bold')
space_1.grid(column=1, row=10)
##-----------------------------------------------------------------------------
ttk.Separator(root, orient='horizontal').grid(row=11, columnspan=200, sticky="ew")
lbl_1 = Label(root, text="Electrical Performance (EBT) - ", font='arial 12 bold')
lbl_1.grid(column=1, row=12) 
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.ebt_num)



labl_12 = Label(root, text="EBT tested quantity")
labl_12.grid(column=1, row=13, sticky=E)
txt_12 = Entry(root, width=30, textvariable=var) 
txt_12.grid(column=2, row=13)
##-----------------------------------------------------------------------------
var = StringVar()
var.set(temp.test_cond)



labl_13 = Label(root, text="Test conditions")
labl_13.grid(column=1, row=14, sticky=E)
txt_13 = Text(root, height = 4, width = 22)
txt_13.grid(column=2, row=14) 
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.test_set_1)

labl_14 = Label(root, text="Test setup")
labl_14.grid(column=1, row=15, sticky=E)
txt_14 = Entry(root, width=30, textvariable=var) 
txt_14.grid(column=2, row=15)
space_2 = Label(root, text=" ", font='arial 12 bold')
space_2.grid(column=1, row=16)
##-----------------------------------------------------------------------------
ttk.Separator(root, orient='horizontal').grid(row=17, columnspan=200, sticky="ew")
lbl_2 = Label(root, text="Frequency Stability over Temperature - ", font='arial 12 bold')
lbl_2.grid(column=1, row=18) 
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.fvt_spec)



labl_15 = Label(root, text="FvT SPEC Limit(ppm)")
labl_15.grid(column=1, row=19, sticky=E)
txt_15 = Entry(root, width=30, textvariable=var) 
txt_15.grid(column=2, row=19)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.temp_range)


labl_16 = Label(root, text="Operating temperature range")
labl_16.grid(column=1, row=20, sticky=E)
txt_16 = Entry(root, width=30, textvariable=var) 
txt_16.grid(column=2, row=20)
txt_16.focus()
##-----------------------------------------------------------------------------
var = StringVar()
var.set(temp.test1_cond)



labl_17 = Label(root, text="Test conditions")
labl_17.grid(column=1, row=21, sticky=E)
txt_17 = Text(root, height = 4, width = 22)
txt_17.grid(column=2, row=21)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.test1_set_1)



labl_18 = Label(root, text="Test setup")
labl_18.grid(column=1, row=22, sticky=E)
txt_18 = Entry(root, width=30, textvariable=var) 
txt_18.grid(column=2, row=22)
space_3 = Label(root, text=" ", font='arial 12 bold')
space_3.grid(column=1, row=23)
##-----------------------------------------------------------------------------
ttk.Separator(root, orient='horizontal').grid(row=24, columnspan=200, sticky="ew")
lbl_3 = Label(root, text="Phase Noise Performance - ", font='arial 12 bold')
lbl_3.grid(column=1, row=25) 
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.plot_num)



labl_19 = Label(root, text="No. of P.N plots")
labl_19.grid(column=1, row=26, sticky=E)
txt_19 = Entry(root, width=30, textvariable=var) 
txt_19.grid(column=2, row=26)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.pn_spec)


labl_21 = Label(root, text="Typ. at '1Hz 10Hz 100Hz 1kHz 10kHz 100kHz 1MHz'")
#labl_21 = Label(root, text="Typ. at '1Hz 10Hz 100Hz 1kHz 10kHz 100kHz 1MHz'")
labl_21.grid(column=1, row=28, sticky=E)
txt_21 = Entry(root, width=30, textvariable=var) 
txt_21.grid(column=2, row=28)
txt_21.focus()
##-----------------------------------------------------------------------------
var = StringVar()
var.set(temp.test2_cond)



labl_22 = Label(root, text="Test conditions")
labl_22.grid(column=1, row=29, sticky=E)
txt_22 = Text(root, height = 4, width = 22)
txt_22.grid(column=2, row=29)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.test2_set_1)



labl_23 = Label(root, text="Test setup")
labl_23.grid(column=1, row=30, sticky=E)
txt_23 = Entry(root, width=30, textvariable=var) 
txt_23.grid(column=2, row=30)
##-----------------------------------------------------------------------------
var = IntVar()
var.set(temp.save_path)



labl_24 = Label(root, text="Name your Sample Data Report")
labl_24.grid(column=1, row=31, sticky=E)
txt_24 = Entry(root, width=30, textvariable=var) 
txt_24.grid(column=2, row=31)
##-----------------------------------------------------------------------------
choices = ['csv', 'txt']
var = StringVar()
var.set(temp.pn_fl_type)


labl_20 = Label(root, text="Data file format")
labl_20.grid(column=3, row=26, sticky=E)
box_20 = OptionMenu(root, var, *choices)
box_20.grid(column=4, row=26, sticky=E)
box_20.config(width=4)
##-----------------------------------------------------------------------------
space_4 = Label(root, text=" ", font='arial 12 bold')
space_4.grid(column=1, row=32)
##-----------------------------------------------------------------------------
ttk.Separator(root, orient='horizontal').grid(row=33, columnspan=200, sticky="ew")
##-----------------------------------------------------------------------------
btn_start = Button(root, text = "GENERATE REPORT", bg='green',fg='white', font='arial 16', command = clicked_start)
btn_start.grid(column = 2, row = 36, sticky=W)

root.mainloop()
#print(temp.pn_spec)

root = tk.Tk()
root.withdraw()  
messagebox.showinfo('Info','Please select the TEMPLATE file(.docx). ')
file_path = filedialog.askopenfilename()
print(file_path)
#file_name = ntpath.basename(file_path)
template = DocxTemplate(file_path) 

#temp = Report_template()

# print(temp.test_cond[0])
# print(temp1.test_cond[0])
# print(temp2.test_cond[0])
# print(temp3.test_cond[0])
context_con = {
    'test_cond_1': temp.test_cond,
    'test1_cond_1': temp.test1_cond,
    'test2_cond_1': temp.test2_cond
    }

#print(context_con)    

context = {
    'device_des': temp.device_des,
    'freq': temp.freq,
    'sup_vol': temp.sup_vol,
    'output_type': temp.output_type,
    'Code': temp.dev_code,
    'day': temp.rev_day,
    'month': temp.rev_month,
    'year': temp.rev_year,
    'n': temp.n,
    'so_num': temp.so_num,
    'd': datetime.datetime.now().strftime('%d'),
    'm': datetime.datetime.now().strftime('%b'),
    'y': datetime.datetime.now().strftime('%Y'),
    'Author': temp.name,
    'test_set_1': temp.test_set_1,
    'test1_set_1': temp.test1_set_1,
    'test2_set_1': temp.test2_set_1
    }

FvT_gen()
save_button = widgets.Button(description = 'Select this Unit data')
ok_button = widgets.Button(description='Proceed')
out = widgets.Output()

temp_data = 0
fvt_limX = [-48, 108]
fvt_limY_pos = [temp.fvt_spec, temp.fvt_spec]
fvt_limY_neg = [-temp.fvt_spec, -temp.fvt_spec]
saved_cnt = 0
#ok_flag = 0

##--------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
def ok_on_click(b):
    with out:
        clear_output()
        #global ok_flag
        #ok_flag = 1
        # Rename column labels in unitNo# sequence  
        unit_qty = len(FvT_gen.res.columns)-1
        col_lst = ["Operating Temperature(degC)"]
        for _ in range(unit_qty):
            col_lst.append("Unit"+ str(_+1))

        print(col_lst)
        FvT_gen.res.columns = col_lst
#---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------

        ## finalised plot generated according to saved Units data
        plt =FvT_gen.res.plot(x = FvT_gen.res.columns[0], y = FvT_gen.res.columns[1:], kind = 'line', figsize = (20, 10), ylabel = 'Norm.Freq(ppm)', title='Sample Frequency Stability', grid = True)
        #plt =FvT_gen.res.plot(x = 0, y = FvT_gen.res.columns[1:], kind = 'line', figsize = (20, 10), ylabel = 'Norm.Freq', title='Sample Frequency Stability', grid = True)
        plt.plot(fvt_limX, fvt_limY_pos, '--r', label = 'FvT Limit+'+ '(' + str(temp.fvt_spec) + 'ppm)')
        plt.plot(fvt_limX, fvt_limY_neg, '--r', label = 'FvT Limit-'+ '(' + str(-temp.fvt_spec) + 'ppm)')

        plt.set_xlim(xmin = temp.temp_range[0], xmax = temp.temp_range[1]) # show preferred x-axis range as per operating temperature
        plt.set_ylim(ymin = -temp.fvt_spec*2.5, ymax = temp.fvt_spec*2.5) # Keep the plot centered within scope

        #plt.xaxis.set_ticks(np.arange(-45, 105, 20))
        #plt.xaxis.set_major_formatter(ticker.FormatStrFormatter('%0.1f'))
        plt.legend()
        fig = plt.get_figure()
        fig.savefig("FvT_output_0" +'.png')
##---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        context['FvT_Plot_0'] = InlineImage(template, 'FvT_output_0.png', width = Mm(240), height= Mm(120))
        #context['FvT_Plot_1'] = InlineImage(template, 'FvT_output_1.png', width = Mm(240), height= Mm(120))
##---------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------------
        if(temp.pn_fl_type == 'csv'):
            PN_gen_csv()
        if(temp.pn_fl_type == 'txt'):
            PN_gen_txt()
    
        context['PN_Plot'] = InlineImage(template, 'PN_output.png', width = Mm(240), height= Mm(120))

        template.render(Merge_context(context, context_con))
        # save_path = simpledialog.askstring(title="Save", prompt="Please type in the file name with .docx : \t\t\t\t\t\t")
        template.save(temp.save_path)
        get_Table_1(temp.save_path)

def save_on_click(b):
    with out:
        #FvT_gen.df2.plot(x = 'Temp(deg)', y = UnitNo, kind = 'line', figsize = (16, 5)).savefig(f'FvT.png')
        #print(FvT_gen.df2[b])
        #print(111)
        clear_output()
        #print(FvT_gen.result[temp_data])
        print("\n.............................................................................................\n")
        print("\nYOUR SAVED UNIT DATA:\n")
        FvT_gen.res[temp_data] = FvT_gen.result[temp_data] ## update selected Unit data into the df3
        print(FvT_gen.res)
        global saved_cnt
        saved_cnt = len(FvT_gen.res.columns)-1
        print("Saved Counter - ", saved_cnt)
        
def multiplot(UnitNo):
    global temp_data ## the global 'temp' variable 
    opts = FvT_gen.result.columns.values
    #plt = FvT_gen.result.plot(x = 0, y = 1, kind = 'line', figsize = (16, 5), xlabel = 'Temperature', ylabel = 'Norm.Freq', title='Sample Frequency Stability', grid = True)
    plt = FvT_gen.result.plot(x = 0, y = UnitNo, kind = 'line', figsize = (16, 5), xlabel = 'Temperature', ylabel = 'Norm.Freq', title='Sample Frequency Stability', grid = True)
    plt.plot(fvt_limX, fvt_limY_pos, '--r', label = 'FvT Limit+')
    plt.plot(fvt_limX, fvt_limY_neg, '--r', label = 'FvT Limit-')
    #plt.xaxis.set_ticks(np.arange(-45, 105, 20))
    #plt.xaxis.set_major_formatter(ticker.FormatStrFormatter('%0.1f'))
    plt.legend()
    save_button.on_click(save_on_click)
    ok_button.on_click(ok_on_click)
    temp_data = locals()['UnitNo'] ## store real-time local 'UnitNo' argument value to global 'temp' using locals()
    #save_button.on_click(save_on_click(UnitNo))
    #ui = widgets.HBox([save_button, out])
#save_button.on_click(save_on_click)


# In[ ]:


cols_2[0] = 1
n = 1
while(n < len(cols_2)):
    cols_2[n] = cols_2[n-1] + 1
    n = n + 1
#print(cols_2)
tuple_Max = max(cols_2)

buttons = widgets.HBox([save_button, ok_button])
display(widgets.HBox([buttons, out]))
interactive_plot = interactive(multiplot, UnitNo= (1, tuple_Max, 1))
output = interactive_plot.children[-1]
#output.layout.height = '350px'
interactive_plot


# In[ ]:


# # Rename column labels in unitNo# sequence  
# unit_qty = len(FvT_gen.res.columns)-1
# col_lst = ["Operating Temperature(degC)"]
# for _ in range(unit_qty):
#     col_lst.append("Unit"+ str(_+1))

# print(col_lst)
# FvT_gen.res.columns = col_lst
# #-------------------------------------------------------------------------------------------------------------------

# ## finalised plot generated according to saved Units data
# plt =FvT_gen.res.plot(x = FvT_gen.res.columns[0], y = FvT_gen.res.columns[1:], kind = 'line', figsize = (20, 10), ylabel = 'Norm.Freq(ppm)', title='Sample Frequency Stability', grid = True)
# #plt =FvT_gen.res.plot(x = 0, y = FvT_gen.res.columns[1:], kind = 'line', figsize = (20, 10), ylabel = 'Norm.Freq', title='Sample Frequency Stability', grid = True)
# plt.plot(fvt_limX, fvt_limY_pos, '--r', label = 'FvT Limit+'+ '(' + str(temp.fvt_spec) + 'ppm)')
# plt.plot(fvt_limX, fvt_limY_neg, '--r', label = 'FvT Limit-'+ '(' + str(-temp.fvt_spec) + 'ppm)')

# plt.set_xlim(xmin = temp.temp_range[0], xmax = temp.temp_range[1]) # show preferred x-axis range as per operating temperature
# plt.set_ylim(ymin = -temp.fvt_spec*2.5, ymax = temp.fvt_spec*2.5) # Keep the plot centered within scope

# #plt.xaxis.set_ticks(np.arange(-45, 105, 20))
# #plt.xaxis.set_major_formatter(ticker.FormatStrFormatter('%0.1f'))
# plt.legend()
# fig = plt.get_figure()
# fig.savefig("FvT_output_0" +'.png')
# ##------------------------------------------------------------------------------------------------------------------
# context['FvT_Plot_0'] = InlineImage(template, 'FvT_output_0.png', width = Mm(240), height= Mm(120))
# #context['FvT_Plot_1'] = InlineImage(template, 'FvT_output_1.png', width = Mm(240), height= Mm(120))
# ##------------------------------------------------------------------------------------------------------------------
# if(temp.pn_fl_type == 'csv'):
#     PN_gen_csv()
# if(temp.pn_fl_type == 'txt'):
#     PN_gen_txt()
    
# context['PN_Plot'] = InlineImage(template, 'PN_output.png', width = Mm(240), height= Mm(120))

# template.render(Merge_context(context, context_con))
# # save_path = simpledialog.askstring(title="Save", prompt="Please type in the file name with .docx : \t\t\t\t\t\t")
# template.save(temp.save_path)
# get_Table_1(temp.save_path)


# In[ ]:




